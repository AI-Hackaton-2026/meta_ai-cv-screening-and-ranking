"""
CV file parsing: PDF and DOCX -> plain text.

Strategy:
  - PDF: try pypdf first (fast, pure-python). Fall back to pdfplumber
    if pypdf returns near-empty text (some PDFs need layout analysis).
  - DOCX: python-docx, concatenate all paragraph and table cell text.
  - Scanned image PDFs are not supported (no OCR); will return empty text
    which the scorer will flag as an error.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx"}


class ParsingError(Exception):
    pass


def validate_upload(filename: str, content: bytes) -> None:
    """Raise ParsingError if the file doesn't meet basic requirements."""
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ParsingError(f"File exceeds 10 MB limit: {filename}")
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ParsingError(f"Unsupported file type '{ext}'. Upload a PDF or DOCX file.")


def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF byte string. Falls back to pdfplumber."""
    # Primary: pypdf (fast)
    try:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if len(text) > 100:
            return text
    except Exception as e:
        logger.warning("pypdf failed, falling back to pdfplumber: %s", e)

    # Fallback: pdfplumber (better with complex layouts)
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
        return text
    except Exception as e:
        logger.error("pdfplumber also failed: %s", e)
        return ""


def parse_docx(content: bytes) -> str:
    """Extract text from a DOCX byte string."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text(filename: str, content: bytes) -> str:
    """
    Entry point: validate the file, then extract and return plain text.
    Raises ParsingError on validation failure.
    Returns an empty string if text extraction yields nothing (e.g. scanned PDF).
    """
    validate_upload(filename, content)

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = parse_pdf(content)
    elif ext == ".docx":
        text = parse_docx(content)
    else:
        raise ParsingError(f"Unsupported extension: {ext}")

    return text
